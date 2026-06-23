#!/usr/bin/env python3
"""
生成报销单Word文档（分组版）
用法: python3 generate_reimbursement_word.py '<json数据>' [图片目录] [输出目录]

分组逻辑：按 level1（一级类型）分组，每组生成一份Word
"""
import sys, json, os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("缺少python-docx，请安装: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="微软雅黑", size=10, bold=False, color=None):
    """设置run字体"""
    font = run.font
    font.name = name
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_image_to_cell(cell, image_path, max_width_cm=2.5, max_height_cm=1.8):
    """在单元格中添加图片，限制大小"""
    if not os.path.exists(image_path):
        return False
    try:
        from PIL import Image as PILImage
        pil_img = PILImage.open(image_path)
        
        # 计算缩放比例
        img_w_cm = pil_img.width / 96 * 2.54  # 假设96dpi
        img_h_cm = pil_img.height / 96 * 2.54
        
        ratio_w = max_width_cm / img_w_cm if img_w_cm > max_width_cm else 1
        ratio_h = max_height_cm / img_h_cm if img_h_cm > max_height_cm else 1
        ratio = min(ratio_w, ratio_h)
        
        new_w = img_w_cm * ratio
        new_h = img_h_cm * ratio
        
        # 清空单元格
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(new_w), height=Cm(new_h))
        return True
    except Exception as e:
        print(f"  嵌入图片失败 {image_path}: {e}")
        return False


def parse_date(date_str):
    """尝试解析日期字符串，返回datetime对象，失败返回None"""
    if not date_str:
        return None
    formats = ['%Y-%m-%d', '%Y/%m/%d', '%m-%d', '%m/%d', '%Y-%m-%d %H:%M:%S', '%Y年%m月%d日']
    for fmt in formats:
        try:
            dt = datetime.strptime(date_str, fmt)
            if fmt in ['%m-%d', '%m/%d']:
                dt = dt.replace(year=datetime.now().year)
            return dt
        except:
            continue
    return None

def sort_by_date(data):
    """按日期从远到近排序（升序），无法解析的日期放最后"""
    def get_date_key(item):
        date_str = item.get('date', '')
        dt = parse_date(date_str)
        if dt:
            return (0, dt)
        return (1, date_str)
    return sorted(data, key=get_date_key)

def create_single_word(data, output_path, image_dir="", doc_title="报销单"):
    """生成单份Word文档"""
    # 按日期从远到近排序
    data = sort_by_date(data)
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading(doc_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, name="微软雅黑", size=16, bold=True, color=(0, 0, 0))
    
    # 创建表格
    headers = ["序号", "日期", "一级类型", "二级类型", "金额(元)", "合计(元)", "凭据截图", "费用说明", "起点", "终点"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头样式
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "4472C4")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 按二级类型分组并写入数据
    total = 0.0
    idx = 0
    i = 0
    while i < len(data):
        item = data[i]
        current_level2 = item.get("level2", "")
        subtotal = 0.0
        sub_items = []
        
        # 收集同二级类型的所有项
        while i < len(data) and data[i].get("level2", "") == current_level2:
            sub_items.append(data[i])
            amount = float(data[i].get("amount", 0) or 0)
            subtotal += amount
            total += amount
            i += 1
        
        # 写入该二级类型的所有明细行
        for sub_item in sub_items:
            idx += 1
            row = table.add_row()
            level1 = sub_item.get("level1", "")
            amount = float(sub_item.get("amount", 0) or 0)
            
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].text = sub_item.get("date", "")
            row.cells[2].text = level1
            row.cells[3].text = sub_item.get("level2", "")
            
            cell = row.cells[4]
            cell.text = f"{amount:.2f}"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            row.cells[5].text = ""
            
            # 凭据截图
            image_name = sub_item.get("image_name", "") or sub_item.get("image", "")
            image_path = ""
            if image_dir and image_name:
                # 尝试多种匹配方式
                possible_paths = []
                # 1. 精确匹配
                possible_paths.append(os.path.join(image_dir, image_name))
                # 2. 去掉路径，只用basename匹配
                basename = os.path.basename(image_name)
                possible_paths.append(os.path.join(image_dir, basename))
                # 3. 模糊匹配 - 遍历目录找包含basename的文件
                for p in Path(image_dir).glob("*"):
                    if p.is_file() and basename in p.name:
                        possible_paths.append(str(p))
                # 4. 反向模糊 - image_name包含在目录文件名中
                for p in Path(image_dir).glob("*"):
                    if p.is_file() and p.name in image_name:
                        possible_paths.append(str(p))
                
                for p in possible_paths:
                    if os.path.exists(p):
                        image_path = p
                        break
            
            cell = row.cells[6]
            if image_path and add_image_to_cell(cell, image_path, max_width_cm=2.5, max_height_cm=1.8):
                pass
            else:
                cell.text = "[图片未找到]"
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 费用说明 — 只有提供了内容才填
            desc = sub_item.get("description", "")
            row.cells[7].text = desc if desc else ""
            
            # 起点终点（仅差旅）- 支持start/end和start_point/end_point
            if level1 == "差旅":
                start = sub_item.get("start_point", "") or sub_item.get("start", "")
                end = sub_item.get("end_point", "") or sub_item.get("end", "")
                row.cells[8].text = start
                row.cells[9].text = end
            else:
                row.cells[8].text = ""
                row.cells[9].text = ""
            
            # 设置字体和垂直居中
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=9)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 插入二级小计行
        subtotal_row = table.add_row()
        subtotal_row.cells[0].text = ""
        subtotal_row.cells[1].text = ""
        subtotal_row.cells[2].text = ""
        subtotal_row.cells[3].text = f"{current_level2}小计"
        
        cell = subtotal_row.cells[4]
        cell.text = f"{subtotal:.2f}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for c in range(5, 10):
            subtotal_row.cells[c].text = ""
        
        for cell in subtotal_row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9, bold=True, color=(128, 128, 128))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, "F2F2F2")
    
    # 总计行
    total_row = table.add_row()
    total_row.cells[0].text = ""
    total_row.cells[1].text = ""
    total_row.cells[2].text = ""
    total_row.cells[3].text = "合计"
    
    cell = total_row.cells[4]
    cell.text = f"{total:.2f}"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    cell = total_row.cells[5]
    cell.text = f"{total:.2f}"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for c in range(6, 10):
        total_row.cells[c].text = ""
    
    for cell in total_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9E2F3")
    
    # 设置列宽
    col_widths = [1.0, 2.0, 2.0, 2.0, 1.8, 1.8, 3.0, 4.5, 2.5, 2.5]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(width)
    
    doc.save(output_path)
    print(f"Word已保存: {output_path}")
    return output_path


def create_grouped_word(data, output_dir, image_dir=""):
    """按一级类型分组，每组生成一份Word"""
    # 先按日期排序（从远到近），再分组
    data = sort_by_date(data)
    groups = {}
    for item in data:
        level1 = item.get("level1", "其他")
        if level1 not in groups:
            groups[level1] = []
        groups[level1].append(item)
    
    output_paths = []
    today = "20260528"
    for level1, items in groups.items():
        safe_name = level1.replace("/", "_").replace("\\", "_")
        output_path = os.path.join(output_dir, f"报销单_{safe_name}_{today}.docx")
        create_single_word(items, output_path, image_dir, doc_title=f"{level1}报销单")
        output_paths.append(output_path)
    
    return output_paths


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 generate_reimbursement_word.py '<json数据>' [图片目录] [输出目录]")
        sys.exit(1)
    
    data = json.loads(sys.argv[1])
    
    if len(sys.argv) >= 4:
        image_dir = sys.argv[2]
        output_dir = sys.argv[3]
        os.makedirs(output_dir, exist_ok=True)
        paths = create_grouped_word(data, output_dir, image_dir)
        print(f"\n生成 {len(paths)} 份报销单:")
        for p in paths:
            print(f"  - {p}")
    else:
        output = sys.argv[2] if len(sys.argv) > 2 else "/root/.openclaw/workspace/报销单.docx"
        create_single_word(data, output)
