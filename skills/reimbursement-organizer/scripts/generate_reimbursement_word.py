#!/usr/bin/env python3
"""
生成报销单Word文档
用法：python3 generate_reimbursement_word.py <json数据文件> <输出目录>

JSON格式：
{
  "records": [
    {
      "date": "2026-05-09",
      "primary_type": "差旅",
      "secondary_type": "交通费",
      "amount": 42.00,
      "start_point": "广东中山市新丰北路",
      "end_point": "聚满分(龙江店)",
      "remarks": "",
      "image_num": 1
    }
  ],
  "images": {
    "1": "/path/to/img_001.jpg",
    "2": "/path/to/img_002.jpg"
  }
}

自动按一级分类拆分，为每个一级分类生成独立的Word文档。
"""

import sys
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from collections import defaultdict

def set_cell_shading(cell, color):
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._element.get_or_add_tcPr().append(shading)

def add_cell_text(cell, text, size=9, bold=False, color=None, align=None):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        cell.paragraphs[0].alignment = align

def add_cell_image(cell, image_path, max_width=Cm(2.5), max_height=Cm(1.8)):
    """在单元格中插入图片，按宽高比例缩放"""
    if not image_path or not os.path.exists(image_path):
        return
    try:
        from PIL import Image
        img = Image.open(image_path)
        img_width, img_height = img.size
        # 计算缩放比例，保持宽高比
        ratio = min(max_width.pt / img_width, max_height.pt / img_height)
        new_width = Pt(img_width * ratio)
        new_height = Pt(img_height * ratio)
        
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run()
        run.add_picture(image_path, width=new_width, height=new_height)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"插入图片失败 {image_path}: {e}")

def generate_single_doc(records, images, doc_title, output_path):
    """生成单份报销单Word文档"""
    if not records:
        print("没有记录数据")
        return False
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run(doc_title)
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 创建表格
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    
    headers = ['序号', '日期', '一级类型', '二级类型', '金额(元)', '合计(元)', '凭据截图', '费用说明', '起点', '终点']
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        add_cell_text(header_cells[i], h, 10, True, (255, 255, 255), WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(header_cells[i], '4472C4')
    
    # 按二级类型分组
    secondary_groups = defaultdict(list)
    for idx, r in enumerate(records, 1):
        secondary_groups[r['secondary_type']].append((idx, r))
    
    # 定义二级类型显示顺序
    type_order = ['交通费', '车费', '油费', '住宿费', '过桥费', '停车费', '餐费', '工作餐', '聚餐费', '活动费', '快递费', '运费', '配送费', '日用品', '办公文具', '通讯费', '招聘费', '办公其他费用', '水果', '牛奶饮料', '蔬菜鲜花', '蛋糕辅料', '其他材料', '兼职工资', '宿舍租金']
    
    # 插入数据行
    for secondary_type in type_order:
        if secondary_type not in secondary_groups:
            continue
        group_records = secondary_groups[secondary_type]
        group_total = sum(r['amount'] for _, r in group_records)
        
        for item_idx, r in group_records:
            row = table.add_row()
            cells = row.cells
            add_cell_text(cells[0], str(item_idx))
            add_cell_text(cells[1], r.get('date', ''))
            add_cell_text(cells[2], r.get('primary_type', ''))
            add_cell_text(cells[3], r.get('secondary_type', ''))
            add_cell_text(cells[4], f"{r['amount']:.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT)
            add_cell_text(cells[5], '')
            # 插入图片 - 凭据截图列
            img_path = images.get(str(item_idx))
            add_cell_image(cells[6], img_path)
            # 费用说明 - 默认空，除非用户有备注
            add_cell_text(cells[7], r.get('remarks', ''))
            add_cell_text(cells[8], r.get('start_point', '') or r.get('start', ''))
            add_cell_text(cells[9], r.get('end_point', '') or r.get('end', ''))
        
        # 小计行
        subtotal_row = table.add_row()
        subtotal_cells = subtotal_row.cells
        for i in range(10):
            add_cell_text(subtotal_cells[i], '')
        add_cell_text(subtotal_cells[3], f'{secondary_type}小计', 9, True)
        add_cell_text(subtotal_cells[4], f'{group_total:.2f}', 9, True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for cell in subtotal_cells:
            set_cell_shading(cell, 'D9D9D9')
    
    # 总计行
    total = sum(r['amount'] for r in records)
    total_row = table.add_row()
    total_cells = total_row.cells
    for i in range(10):
        add_cell_text(total_cells[i], '')
    add_cell_text(total_cells[3], '总计', 9, True, (255, 255, 255))
    add_cell_text(total_cells[4], f'{total:.2f}', 9, True, (255, 255, 255), WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in total_cells:
        set_cell_shading(cell, '4472C4')
    
    doc.save(output_path)
    print(f'保存成功: {output_path}')
    return True

def generate_reimbursement_word(data, output_dir):
    """
    按一级分类自动拆分生成多份报销单Word文档。
    data: 包含 records 和 images 的字典
    output_dir: 输出目录，每份文档按一级分类命名
    """
    records = data.get('records', [])
    images = data.get('images', {})
    
    if not records:
        print("没有记录数据")
        return False
    
    # 按一级分类分组
    primary_groups = defaultdict(list)
    for r in records:
        primary_groups[r['primary_type']].append(r)
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    date_str = datetime.now().strftime('%Y%m%d')
    
    # 为每个一级分类生成独立文档
    for primary_type in sorted(primary_groups.keys()):
        group_records = primary_groups[primary_type]
        
        # 重新编号，并映射图片
        new_records = []
        new_images = {}
        for idx, r in enumerate(group_records, 1):
            new_r = dict(r)
            new_r['id'] = idx
            new_records.append(new_r)
            old_img_num = str(r.get('image_num', ''))
            if old_img_num and old_img_num in images:
                new_images[str(idx)] = images[old_img_num]
        
        doc_title = f'{primary_type}报销单'
        output_path = os.path.join(output_dir, f'报销单_{primary_type}_{date_str}.docx')
        
        generate_single_doc(new_records, new_images, doc_title, output_path)
    
    print(f"\n共生成 {len(primary_groups)} 份报销单，保存目录: {output_dir}")
    return True

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python3 generate_reimbursement_word.py <json文件> <输出目录>")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_dir = sys.argv[2]
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    generate_reimbursement_word(data, output_dir)
