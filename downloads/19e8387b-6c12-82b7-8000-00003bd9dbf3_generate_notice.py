#!/usr/bin/env python3
"""
Generate training summary notice Word document.
Usage: python3 generate_notice.py --config config.json --output output.docx
"""

import json
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcPr.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')

def generate_notice(config_path, output_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get('title', '培训总结通知'))
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Greeting
    if data.get('greeting'):
        p = doc.add_paragraph(data['greeting'])
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    # Overview
    if data.get('overview'):
        p = doc.add_paragraph(data['overview'])
        p.paragraph_format.space_after = Pt(12)
    
    # Section 1: Participants
    if data.get('participants'):
        doc.add_paragraph().add_run('一、参训人员').bold = True
        
        participants = data['participants']
        if participants:
            headers = list(participants[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # Data rows
            for person in participants:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(person.get(header, ''))
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            doc.add_paragraph()  # spacing
    
    # Section 2: Training modules
    if data.get('modules'):
        doc.add_paragraph().add_run('二、培训回顾').bold = True
        
        modules = data['modules']
        if modules:
            headers = list(modules[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            for module in modules:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(module.get(header, ''))
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            doc.add_paragraph()
    
    # Section 3: Awards
    if data.get('team_award') or data.get('individual_awards'):
        doc.add_paragraph().add_run('三、表彰').bold = True
        
        if data.get('team_award'):
            team = data['team_award']
            p = doc.add_paragraph()
            p.add_run(f"精英团队：{team.get('team_name', '')}").bold = True
            for member in team.get('members', []):
                doc.add_paragraph(member, style='List Bullet')
        
        if data.get('individual_awards'):
            p = doc.add_paragraph()
            p.add_run('优秀学员').bold = True
            p.paragraph_format.space_before = Pt(6)
            for name in data['individual_awards']:
                doc.add_paragraph(name, style='List Bullet')
        
        doc.add_paragraph('再次祝贺以上获奖团队和学员！')
        doc.add_paragraph()
    
    # Section 4: Reminders
    has_reminders = data.get('reminders_to_trainees') or data.get('reminders_to_mentors')
    if has_reminders:
        doc.add_paragraph().add_run('四、重要叮嘱').bold = True
        
        if data.get('reminders_to_trainees'):
            p = doc.add_paragraph()
            p.add_run('致所有学员：').bold = True
            p = doc.add_paragraph(data['reminders_to_trainees'])
            p.paragraph_format.first_line_indent = Inches(0.3)
        
        if data.get('reminders_to_mentors'):
            p = doc.add_paragraph()
            p.add_run('致所有带训师傅：').bold = True
            p.paragraph_format.space_before = Pt(6)
            p = doc.add_paragraph(data['reminders_to_mentors'])
            p.paragraph_format.first_line_indent = Inches(0.3)
        
        doc.add_paragraph()
    
    # Closing
    if data.get('closing'):
        p = doc.add_paragraph(data['closing'])
        p.paragraph_format.space_after = Pt(12)
    
    # Contact
    if data.get('contact'):
        p = doc.add_paragraph(data['contact'])
        p.paragraph_format.space_after = Pt(12)
    
    # Department and Date (right aligned)
    if data.get('department') or data.get('date'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if data.get('department'):
            p.add_run(data['department'])
            if data.get('date'):
                p.add_run('\n')
        if data.get('date'):
            p.add_run(data['date'])
    
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate training summary notice')
    parser.add_argument('--config', required=True, help='Path to JSON config file')
    parser.add_argument('--output', required=True, help='Output docx file path')
    args = parser.parse_args()
    
    generate_notice(args.config, args.output)
